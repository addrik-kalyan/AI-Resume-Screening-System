"""
utils/resume_parser.py
-----------------------
Module 3: Resume Parser
Extracts raw text from an uploaded resume file (PDF or DOCX).

Technology: PDFPlumber (PDF), python-docx (DOCX)
"""

import os
import pdfplumber
import docx


def extract_text(file_path):
    """
    Extract plain text from a resume file.
    Supports .pdf and .docx
    Returns: str (extracted text, lowercase-safe for downstream NLP)
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported resume file type: {ext}. Use PDF or DOCX.")


def _extract_from_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def _extract_from_docx(file_path):
    document = docx.Document(file_path)
    text = "\n".join(p.text for p in document.paragraphs)
    # also capture text inside tables (some resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text.strip()
